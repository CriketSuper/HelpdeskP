import re
from copy import deepcopy
from functools import lru_cache
from html import escape
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree as ET

from django.conf import settings
from django.utils import timezone
from docx import Document as WordDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate
from pymorphy3 import MorphAnalyzer
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph as PdfParagraph
from reportlab.platypus import SimpleDocTemplate, Spacer, Table, TableStyle

from .models import UserProfile
from .rich_text import rich_text_has_text, rich_text_to_plain_text, sanitize_rich_text


TICKET_TEMPLATE_PATH = Path(settings.BASE_DIR) / "static" / "docs" / "ticket_template.docx"
CONTENT_PLACEHOLDER = "__HELPDESK_TICKET_CONTENT__"
MONTH_NAMES = {
    1: "января",
    2: "февраля",
    3: "марта",
    4: "апреля",
    5: "мая",
    6: "июня",
    7: "июля",
    8: "августа",
    9: "сентября",
    10: "октября",
    11: "ноября",
    12: "декабря",
}
CYRILLIC_RE = re.compile(r"^[А-Яа-яЁё-]+$")
ALIGNMENTS = {"left", "center", "right"}


def _local_name(tag):
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def _append_segment(segments, text, *, bold=False, italic=False, underline=False):
    if not text:
        return
    if segments and all(
        (
            segments[-1]["bold"] == bold,
            segments[-1]["italic"] == italic,
            segments[-1]["underline"] == underline,
        )
    ):
        segments[-1]["text"] += text
        return
    segments.append(
        {
            "text": text,
            "bold": bold,
            "italic": italic,
            "underline": underline,
        }
    )


def _get_alignment(node):
    alignment = (node.attrib.get("align") or "").strip().lower()
    return alignment if alignment in ALIGNMENTS else None


def _sanitize_content_html(value):
    sanitized = sanitize_rich_text(value)
    if not rich_text_has_text(sanitized):
        return ""
    return sanitized.replace("&nbsp;", "&#160;")


def _append_inline_segments_from_node(node, segments, state):
    text = node.text or ""
    if text:
        _append_segment(segments, text, **state)

    for child in node:
        tag = _local_name(child.tag).lower()
        child_state = deepcopy(state)
        if tag in {"strong", "b"}:
            child_state["bold"] = True
        elif tag in {"em", "i"}:
            child_state["italic"] = True
        elif tag == "u":
            child_state["underline"] = True

        if tag == "br":
            _append_segment(segments, "\n", **state)
        elif tag in {"strong", "b", "em", "i", "u", "a", "span"}:
            _append_inline_segments_from_node(child, segments, child_state)
        else:
            _append_inline_segments_from_node(child, segments, child_state)

        tail = child.tail or ""
        if tail:
            _append_segment(segments, tail, **state)


def _build_paragraph_block(node, *, block_type="paragraph", prefix=""):
    segments = []
    if prefix:
        _append_segment(segments, prefix)
    _append_inline_segments_from_node(
        node,
        segments,
        {"bold": False, "italic": False, "underline": False},
    )
    if not any(segment["text"].strip() for segment in segments):
        return None
    return {
        "type": block_type,
        "align": _get_alignment(node),
        "segments": segments,
    }


def _parse_list_block(node):
    blocks = []
    ordered_index = 0
    for child in node:
        if _local_name(child.tag).lower() != "li":
            continue
        ordered_index += 1
        prefix = f"{ordered_index}. " if _local_name(node.tag).lower() == "ol" else "• "
        block = _build_paragraph_block(child, block_type="list_item", prefix=prefix)
        if block:
            blocks.append(block)
    return blocks


def _parse_table_cell(cell):
    cell_blocks = []
    if (cell.text or "").strip():
        cell_blocks.append(
            {
                "type": "paragraph",
                "align": _get_alignment(cell),
                "segments": [
                    {
                        "text": cell.text or "",
                        "bold": False,
                        "italic": False,
                        "underline": False,
                    }
                ],
            }
        )

    for child in cell:
        tag = _local_name(child.tag).lower()
        if tag == "p":
            block = _build_paragraph_block(child)
            if block:
                cell_blocks.append(block)
        elif tag == "blockquote":
            block = _build_paragraph_block(child, block_type="blockquote")
            if block:
                cell_blocks.append(block)
        elif tag in {"ul", "ol"}:
            cell_blocks.extend(_parse_list_block(child))
        elif tag in {"strong", "b", "em", "i", "u", "a", "span"}:
            wrapper = ET.Element("p")
            wrapper.append(deepcopy(child))
            block = _build_paragraph_block(wrapper)
            if block:
                cell_blocks.append(block)

        if (child.tail or "").strip():
            tail_wrapper = ET.Element("p")
            tail_wrapper.text = child.tail
            block = _build_paragraph_block(tail_wrapper)
            if block:
                cell_blocks.append(block)

    if not cell_blocks:
        cell_blocks.append(
            {
                "type": "paragraph",
                "align": _get_alignment(cell),
                "segments": [{"text": "", "bold": False, "italic": False, "underline": False}],
            }
        )

    return {
        "header": _local_name(cell.tag).lower() == "th",
        "align": _get_alignment(cell),
        "blocks": cell_blocks,
    }


def _parse_table_block(node):
    rows = []
    row_nodes = []
    for child in node:
        tag = _local_name(child.tag).lower()
        if tag == "tr":
            row_nodes.append(child)
        elif tag in {"thead", "tbody"}:
            row_nodes.extend(grand for grand in child if _local_name(grand.tag).lower() == "tr")

    for row_node in row_nodes:
        row = []
        for cell in row_node:
            if _local_name(cell.tag).lower() in {"th", "td"}:
                row.append(_parse_table_cell(cell))
        if row:
            rows.append(row)

    if not rows:
        return None

    return {
        "type": "table",
        "align": _get_alignment(node),
        "rows": rows,
    }


def _parse_container_blocks(element):
    blocks = []

    if (element.text or "").strip():
        intro = ET.Element("p")
        intro.text = element.text
        block = _build_paragraph_block(intro)
        if block:
            blocks.append(block)

    for child in element:
        tag = _local_name(child.tag).lower()
        if tag == "p":
            block = _build_paragraph_block(child)
            if block:
                blocks.append(block)
        elif tag == "blockquote":
            block = _build_paragraph_block(child, block_type="blockquote")
            if block:
                blocks.append(block)
        elif tag in {"ul", "ol"}:
            blocks.extend(_parse_list_block(child))
        elif tag == "table":
            block = _parse_table_block(child)
            if block:
                blocks.append(block)
        else:
            wrapper = ET.Element("p")
            wrapper.append(deepcopy(child))
            block = _build_paragraph_block(wrapper)
            if block:
                blocks.append(block)

        if (child.tail or "").strip():
            tail_block = ET.Element("p")
            tail_block.text = child.tail
            block = _build_paragraph_block(tail_block)
            if block:
                blocks.append(block)

    return blocks


def _parse_rich_text_blocks(value):
    sanitized = _sanitize_content_html(value)
    if not sanitized:
        return []
    root = ET.fromstring(f"<root>{sanitized}</root>")
    return _parse_container_blocks(root)


def _get_user_source_name(user):
    if user is None:
        return ""
    try:
        return user.profile.verbose_name
    except UserProfile.DoesNotExist:
        return user.get_username()


def _get_user_position(user):
    if user is None:
        return ""
    try:
        return (user.profile.position or "").strip()
    except UserProfile.DoesNotExist:
        return ""


@lru_cache(maxsize=1)
def _morph():
    return MorphAnalyzer()


def _split_person_name(full_name):
    parts = [part for part in (full_name or "").split() if part]
    surname = parts[0] if parts else ""
    name = parts[1] if len(parts) > 1 else ""
    patronymic = parts[2] if len(parts) > 2 else ""
    remainder = parts[3:] if len(parts) > 3 else []
    return surname, name, patronymic, remainder


def _build_initials(name, patronymic):
    initials = []
    if name:
        initials.append(f"{name[0]}.")
    if patronymic:
        initials.append(f"{patronymic[0]}.")
    return "".join(initials)


def _apply_casing(source, value):
    if source.isupper():
        return value.upper()
    if source.istitle():
        return value.capitalize()
    return value


def _pick_parse(word, allowed_pos=None, require_grammeme=None, preferred_grammemes=None):
    allowed_pos = set(allowed_pos or [])
    preferred_grammemes = set(preferred_grammemes or [])
    parses = _morph().parse(word)

    for parse in parses:
        if require_grammeme and require_grammeme not in parse.tag:
            continue
        if allowed_pos and parse.tag.POS not in allowed_pos:
            continue
        if preferred_grammemes and not preferred_grammemes.issubset(set(parse.tag.grammemes)):
            continue
        return parse

    for parse in parses:
        if allowed_pos and parse.tag.POS not in allowed_pos:
            continue
        return parse

    return parses[0] if parses else None


def _inflect_word(
    word,
    grammatical_case,
    *,
    allowed_pos=None,
    require_grammeme=None,
    preferred_grammemes=None,
):
    if not word or not CYRILLIC_RE.match(word):
        return word
    parse = _pick_parse(
        word,
        allowed_pos=allowed_pos,
        require_grammeme=require_grammeme,
        preferred_grammemes=preferred_grammemes,
    )
    if parse is None:
        return word
    inflected = parse.inflect({grammatical_case})
    if inflected is None:
        return word
    return _apply_casing(word, inflected.word)


def _infer_person_gender(name, patronymic):
    patronymic_normalized = (patronymic or "").strip().lower()
    if patronymic_normalized.endswith("вна"):
        return "femn"
    if patronymic_normalized.endswith("ич"):
        return "masc"

    name_normalized = (name or "").strip()
    if not name_normalized or not CYRILLIC_RE.match(name_normalized):
        return None

    for parse in _morph().parse(name_normalized):
        grammemes = set(parse.tag.grammemes)
        if "Name" not in grammemes:
            continue
        if "femn" in grammemes:
            return "femn"
        if "masc" in grammemes:
            return "masc"
    return None


def _inflect_position(position, grammatical_case):
    words = [word for word in (position or "").split() if word]
    if not words:
        return ""

    parsed_words = [_pick_parse(word) for word in words]
    head_index = None
    for index, parse in enumerate(parsed_words):
        if parse and parse.tag.POS == "NOUN":
            head_index = index
            break

    if head_index is None:
        return " ".join(_inflect_word(word, grammatical_case) for word in words)

    inflected_words = []
    for index, word in enumerate(words):
        parse = parsed_words[index]
        if parse is None:
            inflected_words.append(word)
            continue
        if index < head_index and parse.tag.POS in {"ADJF", "PRTF"}:
            inflected_words.append(_inflect_word(word, grammatical_case, allowed_pos={"ADJF", "PRTF"}))
        elif index == head_index:
            inflected_words.append(_inflect_word(word, grammatical_case, allowed_pos={"NOUN"}))
        else:
            inflected_words.append(word)
    return " ".join(inflected_words)


def _format_person_for_document(user, grammatical_case, *, initials_first=False):
    full_name = _get_user_source_name(user)
    surname, name, patronymic, remainder = _split_person_name(full_name)
    if not surname:
        return full_name

    gender = _infer_person_gender(name, patronymic)
    inflected_surname = _inflect_word(
        surname,
        grammatical_case,
        allowed_pos={"NOUN", "ADJF"},
        require_grammeme="Surn",
        preferred_grammemes={gender} if gender else None,
    )
    initials = _build_initials(name, patronymic)
    surname_tail = " ".join([inflected_surname, *remainder]).strip()

    if not initials:
        return surname_tail
    if initials_first:
        return f"{initials} {surname_tail}".strip()
    return f"{surname_tail} {initials}".strip()


def _format_ticket_document_date(ticket):
    source_date = timezone.localtime(ticket.created_at) if ticket.created_at else timezone.localtime()
    return f"«{source_date.day:02d}» {MONTH_NAMES[source_date.month]} {source_date.year} г."


def build_ticket_document_context(ticket):
    return {
        "technician_position_dative": _inflect_position(_get_user_position(ticket.technician), "datv"),
        "organization_name": getattr(settings, "ORGANIZATION_NAME", "Helpdesk"),
        "technician_name_dative": _format_person_for_document(ticket.technician, "datv", initials_first=True),
        "author_position_accusative": _inflect_position(_get_user_position(ticket.created_by), "accs"),
        "author_name_accusative": _format_person_for_document(ticket.created_by, "accs", initials_first=False),
        "ticket_content_plain": rich_text_to_plain_text(ticket.content),
        "ticket_content_rich": CONTENT_PLACEHOLDER,
        "ticket_date": _format_ticket_document_date(ticket),
    }


def build_ticket_docx_filename(ticket):
    return f"ticket_{ticket.pk}.docx"


def build_ticket_pdf_filename(ticket):
    return f"ticket_{ticket.pk}.pdf"


def _docx_alignment_value(alignment):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(alignment)


def _clear_docx_paragraph(paragraph):
    paragraph_element = paragraph._element
    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)


def _remove_docx_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def _insert_docx_paragraph_after(element, parent):
    paragraph_element = OxmlElement("w:p")
    element.addnext(paragraph_element)
    return Paragraph(paragraph_element, parent)


def _apply_docx_run_style(run, segment):
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = segment["bold"]
    run.italic = segment["italic"]
    run.underline = segment["underline"]
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "Times New Roman")
    fonts.set(qn("w:cs"), "Times New Roman")


def _populate_docx_paragraph(paragraph, block, *, default_alignment="left", first_line_indent=True):
    _clear_docx_paragraph(paragraph)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = None

    alignment = block.get("align") or default_alignment
    paragraph.alignment = _docx_alignment_value(alignment)

    if block["type"] == "blockquote":
        paragraph.paragraph_format.left_indent = Mm(8)
    elif first_line_indent and alignment in {None, "left"} and block["type"] in {"paragraph", "list_item"}:
        paragraph.paragraph_format.first_line_indent = Mm(12.5)

    for segment in block["segments"]:
        parts = segment["text"].split("\n")
        for index, part in enumerate(parts):
            run = paragraph.add_run(part)
            _apply_docx_run_style(run, segment)
            if index < len(parts) - 1:
                run.add_break(WD_BREAK.LINE)


def _populate_docx_cell(cell, cell_info):
    while len(cell.paragraphs) > 1:
        paragraph = cell.paragraphs[-1]
        paragraph._element.getparent().remove(paragraph._element)

    default_paragraph = cell.paragraphs[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    for block_index, block in enumerate(cell_info["blocks"]):
        paragraph = default_paragraph if block_index == 0 else cell.add_paragraph()
        _populate_docx_paragraph(
            paragraph,
            block,
            default_alignment=cell_info.get("align") or block.get("align") or "left",
            first_line_indent=False,
        )


def _build_docx_table(document, table_block):
    row_count = len(table_block["rows"])
    column_count = max(len(row) for row in table_block["rows"])
    table = document.add_table(rows=row_count, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    for row_index, row in enumerate(table_block["rows"]):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            if column_index >= len(row):
                _populate_docx_cell(
                    cell,
                    {
                        "align": "left",
                        "blocks": [{"type": "paragraph", "align": "left", "segments": [{"text": "", "bold": False, "italic": False, "underline": False}]}],
                    },
                )
                continue
            _populate_docx_cell(cell, row[column_index])
    return table


def _inject_docx_content(document, blocks):
    placeholder_paragraph = next(
        (paragraph for paragraph in document.paragraphs if CONTENT_PLACEHOLDER in paragraph.text),
        None,
    )
    if placeholder_paragraph is None:
        title_index = next(
            (
                index
                for index, paragraph in enumerate(document.paragraphs)
                if paragraph.text.strip() == "СЛУЖЕБНАЯ ЗАПИСКА"
            ),
            None,
        )
        if title_index is not None:
            for paragraph in document.paragraphs[title_index + 1 :]:
                if "\t" in paragraph.text or (
                    paragraph.text.strip() == ""
                    and any("\t" in run.text for run in paragraph.runs)
                ):
                    placeholder_paragraph = paragraph
                    break
    if placeholder_paragraph is None:
        return

    if not blocks:
        _clear_docx_paragraph(placeholder_paragraph)
        return

    parent = placeholder_paragraph._parent
    last_element = placeholder_paragraph._element
    placeholder_used = False
    placeholder_removed = False

    for block_index, block in enumerate(blocks):
        if block["type"] == "table":
            table = _build_docx_table(document, block)
            last_element.addnext(table._tbl)
            last_element = table._tbl
            if block_index == 0:
                _remove_docx_paragraph(placeholder_paragraph)
                placeholder_removed = True
        else:
            if block_index == 0 and not placeholder_removed:
                paragraph = placeholder_paragraph
                placeholder_used = True
            else:
                paragraph = _insert_docx_paragraph_after(last_element, parent)
            _populate_docx_paragraph(paragraph, block)
            last_element = paragraph._element

    if not placeholder_used and not placeholder_removed:
        _remove_docx_paragraph(placeholder_paragraph)


def render_ticket_docx(ticket):
    template = DocxTemplate(str(TICKET_TEMPLATE_PATH))
    template.render(build_ticket_document_context(ticket))
    rendered = BytesIO()
    template.save(rendered)
    rendered.seek(0)

    document = WordDocument(rendered)
    _inject_docx_content(document, _parse_rich_text_blocks(ticket.content))

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


def _find_font_path(candidates):
    for candidate in candidates:
        font_path = Path(candidate)
        if font_path.exists():
            return font_path
    return None


@lru_cache(maxsize=1)
def _register_pdf_fonts():
    regular_font_path = _find_font_path(
        [
            Path(settings.BASE_DIR) / "static" / "fonts" / "times.ttf",
            Path("C:/Windows/Fonts/times.ttf"),
            Path("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf"),
            Path("/usr/share/fonts/truetype/liberation2/LiberationSerif-Regular.ttf"),
            Path("/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf"),
        ]
    )
    bold_font_path = _find_font_path(
        [
            Path(settings.BASE_DIR) / "static" / "fonts" / "timesbd.ttf",
            Path("C:/Windows/Fonts/timesbd.ttf"),
            Path("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf"),
            Path("/usr/share/fonts/truetype/liberation2/LiberationSerif-Bold.ttf"),
            Path("/usr/share/fonts/truetype/liberation/LiberationSerif-Bold.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf"),
        ]
    )
    italic_font_path = _find_font_path(
        [
            Path(settings.BASE_DIR) / "static" / "fonts" / "timesi.ttf",
            Path("C:/Windows/Fonts/timesi.ttf"),
            Path("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Italic.ttf"),
            Path("/usr/share/fonts/truetype/liberation2/LiberationSerif-Italic.ttf"),
            Path("/usr/share/fonts/truetype/liberation/LiberationSerif-Italic.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif-Italic.ttf"),
        ]
    )
    bold_italic_font_path = _find_font_path(
        [
            Path(settings.BASE_DIR) / "static" / "fonts" / "timesbi.ttf",
            Path("C:/Windows/Fonts/timesbi.ttf"),
            Path("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold_Italic.ttf"),
            Path("/usr/share/fonts/truetype/liberation2/LiberationSerif-BoldItalic.ttf"),
            Path("/usr/share/fonts/truetype/liberation/LiberationSerif-BoldItalic.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif-BoldItalic.ttf"),
        ]
    )

    if regular_font_path is None:
        raise FileNotFoundError("No serif font with Cyrillic support was found for PDF export.")

    regular_name = "TicketDocRegular"
    bold_name = "TicketDocBold"
    italic_name = "TicketDocItalic"
    bold_italic_name = "TicketDocBoldItalic"

    if regular_name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(regular_name, str(regular_font_path)))
    if bold_font_path is not None and bold_name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(bold_name, str(bold_font_path)))
    if italic_font_path is not None and italic_name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(italic_name, str(italic_font_path)))
    if bold_italic_font_path is not None and bold_italic_name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(bold_italic_name, str(bold_italic_font_path)))

    if bold_font_path is None:
        bold_name = regular_name
    if italic_font_path is None:
        italic_name = regular_name
    if bold_italic_font_path is None:
        bold_italic_name = bold_name if italic_name == regular_name else italic_name

    return regular_name, bold_name, italic_name, bold_italic_name


def _segment_to_pdf_markup(segment, fonts):
    regular_font, bold_font, italic_font, bold_italic_font = fonts
    font_name = regular_font
    if segment["bold"] and segment["italic"]:
        font_name = bold_italic_font
    elif segment["bold"]:
        font_name = bold_font
    elif segment["italic"]:
        font_name = italic_font

    text = escape(segment["text"]).replace("\n", "<br/>")
    if segment["underline"]:
        text = f"<u>{text}</u>"
    return f'<font name="{font_name}">{text}</font>'


def _pdf_alignment_value(alignment):
    return {
        "left": TA_LEFT,
        "center": TA_CENTER,
        "right": TA_RIGHT,
    }.get(alignment or "left", TA_LEFT)


def _build_pdf_paragraph_style(base_style, name, *, alignment="left", quote=False, first_line_indent=True):
    style = ParagraphStyle(
        name,
        parent=base_style,
        alignment=_pdf_alignment_value(alignment),
    )
    if quote:
        style.leftIndent = 8 * mm
        style.firstLineIndent = 0
    elif first_line_indent and alignment in {None, "left"}:
        style.firstLineIndent = 12.5 * mm
    else:
        style.firstLineIndent = 0
    return style


def _build_pdf_paragraph_for_block(block, styles, fonts):
    markup = "".join(_segment_to_pdf_markup(segment, fonts) for segment in block["segments"])
    style_map = {
        ("paragraph", "left"): styles["content_left"],
        ("paragraph", "center"): styles["content_center"],
        ("paragraph", "right"): styles["content_right"],
        ("list_item", "left"): styles["content_left"],
        ("list_item", "center"): styles["content_center"],
        ("list_item", "right"): styles["content_right"],
        ("blockquote", "left"): styles["quote_left"],
        ("blockquote", "center"): styles["quote_center"],
        ("blockquote", "right"): styles["quote_right"],
    }
    style = style_map.get((block["type"], block.get("align") or "left"), styles["content_left"])
    return PdfParagraph(markup, style)


def _build_pdf_table_flowable(table_block, styles, fonts):
    table_data = []
    header_cells = []

    for row_index, row in enumerate(table_block["rows"]):
        rendered_row = []
        for column_index, cell in enumerate(row):
            cell_markup_parts = []
            for block_index, block in enumerate(cell["blocks"]):
                if block_index:
                    cell_markup_parts.append("<br/><br/>")
                cell_markup_parts.append("".join(_segment_to_pdf_markup(segment, fonts) for segment in block["segments"]))
            paragraph = PdfParagraph(
                "".join(cell_markup_parts),
                {
                    "left": styles["table_left"],
                    "center": styles["table_center"],
                    "right": styles["table_right"],
                }.get(cell.get("align") or "left", styles["table_left"]),
            )
            rendered_row.append(paragraph)
            if cell["header"]:
                header_cells.append((column_index, row_index))
        table_data.append(rendered_row)

    flowable = Table(table_data, repeatRows=1 if header_cells else 0)
    table_style_commands = [
        ("GRID", (0, 0), (-1, -1), 0.6, "#cbd5e1"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    for column_index, row_index in header_cells:
        table_style_commands.append(("BACKGROUND", (column_index, row_index), (column_index, row_index), "#f8fafc"))
    flowable.setStyle(TableStyle(table_style_commands))
    return flowable


def render_ticket_pdf(ticket):
    fonts = _register_pdf_fonts()
    context = build_ticket_document_context(ticket)

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=30 * mm,
        rightMargin=15 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    regular_font_name = fonts[0]
    bold_font_name = fonts[1]
    styles = getSampleStyleSheet()
    regular_style = ParagraphStyle(
        "TicketRegular",
        parent=styles["Normal"],
        fontName=regular_font_name,
        fontSize=14,
        leading=18,
    )
    styles_map = {
        "regular": regular_style,
        "right": ParagraphStyle("TicketRight", parent=regular_style, alignment=TA_RIGHT),
        "title": ParagraphStyle(
            "TicketTitle",
            parent=regular_style,
            fontName=bold_font_name,
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "content_left": _build_pdf_paragraph_style(regular_style, "TicketContentLeft", alignment="left"),
        "content_center": _build_pdf_paragraph_style(regular_style, "TicketContentCenter", alignment="center", first_line_indent=False),
        "content_right": _build_pdf_paragraph_style(regular_style, "TicketContentRight", alignment="right", first_line_indent=False),
        "quote_left": _build_pdf_paragraph_style(regular_style, "TicketQuoteLeft", alignment="left", quote=True),
        "quote_center": _build_pdf_paragraph_style(regular_style, "TicketQuoteCenter", alignment="center", quote=True),
        "quote_right": _build_pdf_paragraph_style(regular_style, "TicketQuoteRight", alignment="right", quote=True),
        "table_left": ParagraphStyle("TicketTableLeft", parent=regular_style, alignment=TA_LEFT, leading=17),
        "table_center": ParagraphStyle("TicketTableCenter", parent=regular_style, alignment=TA_CENTER, leading=17),
        "table_right": ParagraphStyle("TicketTableRight", parent=regular_style, alignment=TA_RIGHT, leading=17),
    }

    story = []
    top_right_block = [
        PdfParagraph(context["technician_position_dative"], styles_map["right"]),
        PdfParagraph(context["organization_name"], styles_map["right"]),
        PdfParagraph(context["technician_name_dative"], styles_map["right"]),
        Spacer(1, 4),
        PdfParagraph(context["author_position_accusative"], styles_map["right"]),
        PdfParagraph(context["author_name_accusative"], styles_map["right"]),
    ]
    header_table = Table([["", top_right_block]], colWidths=[95 * mm, 70 * mm])
    header_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(header_table)
    story.append(Spacer(1, 14))
    story.append(PdfParagraph("СЛУЖЕБНАЯ ЗАПИСКА", styles_map["title"]))
    story.append(Spacer(1, 8))

    blocks = _parse_rich_text_blocks(ticket.content)
    if not blocks:
        blocks = [
            {
                "type": "paragraph",
                "align": "left",
                "segments": [{"text": rich_text_to_plain_text(ticket.content), "bold": False, "italic": False, "underline": False}],
            }
        ]

    for block in blocks:
        if block["type"] == "table":
            story.append(_build_pdf_table_flowable(block, styles_map, fonts))
        else:
            story.append(_build_pdf_paragraph_for_block(block, styles_map, fonts))
        story.append(Spacer(1, 6))
    if len(story) > 4:
        story.pop()

    story.append(Spacer(1, 28))
    footer_table = Table(
        [[PdfParagraph(context["ticket_date"], styles_map["regular"]), PdfParagraph("________________", styles_map["right"])]],
        colWidths=[80 * mm, 85 * mm],
    )
    footer_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(footer_table)

    document.build(story)
    buffer.seek(0)
    return buffer.getvalue()
